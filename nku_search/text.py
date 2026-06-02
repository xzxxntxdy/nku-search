from __future__ import annotations

from collections import Counter
from html import unescape
import io
import math
import re
from urllib.parse import unquote, urlparse
import warnings
import zipfile
from typing import Iterable

TOKEN_RE = re.compile(r"[\u4e00-\u9fff]+|[A-Za-z0-9_]+", re.UNICODE)
CJK_RE = re.compile(r"[\u4e00-\u9fff]+")
SPACE_RE = re.compile(r"\s+")
CONTROL_CHAR_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
PRIVATE_USE_RE = re.compile(r"[\ue000-\uf8ff]")
PERCENT_ESCAPE_RE = re.compile(r"%[0-9A-Fa-f]{2}")
DOCUMENT_EXTENSION_RE = re.compile(r"\.(?:pdf|doc|docx|xls|xlsx|ppt|pptx|txt)$", re.IGNORECASE)
UUIDISH_RE = re.compile(r"^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$", re.IGNORECASE)
MOJIBAKE_MARKERS = ("\ufffd", "\u951f", "\u00c3", "\u00c2", "\u00e5", "\u00e6", "\u00e7", "\u00e9")
BINARY_TEXT_MARKERS = ("[content types].xml", "[content_types].xml", "jfif", "\ufffd")
DOCUMENT_FALLBACK_TITLES = {
    "pdf": "PDF document",
    "doc": "DOC document",
    "docx": "DOCX document",
    "xls": "XLS document",
    "xlsx": "XLSX document",
    "ppt": "PPT document",
    "pptx": "PPTX document",
    "txt": "TXT document",
}


def normalize_space(value: str) -> str:
    text = unescape(value or "")
    text = CONTROL_CHAR_RE.sub(" ", text)
    text = PRIVATE_USE_RE.sub(" ", text)
    return SPACE_RE.sub(" ", text).strip()


def _cjk_char_count(value: str) -> int:
    return sum(1 for char in value if "\u4e00" <= char <= "\u9fff")


def is_probably_binary_text(value: str) -> bool:
    raw = str(value or "")
    if raw.count("\x00") >= 3:
        return True
    cleaned = normalize_space(raw).strip()
    if not cleaned:
        return False
    if cleaned.startswith("\u0871") or "\u0871" in cleaned[:40]:
        return True
    lowered = cleaned[:240].lower()
    if lowered.startswith("pk") and any(marker in lowered for marker in BINARY_TEXT_MARKERS):
        return True
    if any(marker in lowered for marker in BINARY_TEXT_MARKERS):
        return _cjk_char_count(cleaned) == 0
    printable = sum(1 for char in cleaned if char.isprintable())
    searchable = sum(1 for char in cleaned if char.isalnum() or "\u4e00" <= char <= "\u9fff")
    if printable / max(len(cleaned), 1) < 0.85:
        return True
    return len(cleaned) >= 30 and searchable / max(len(cleaned), 1) < 0.12


def clean_document_text(text: str, filetype: str = "") -> str:
    cleaned = normalize_space(text)
    if filetype.lower() in {"ppt", "pptx"} and is_probably_binary_text(cleaned):
        return ""
    return cleaned


def _has_title_signal(value: str) -> bool:
    if not value or is_probably_binary_text(value):
        return False
    cjk = _cjk_char_count(value)
    latin = sum(1 for char in value if char.isalpha() and not ("\u4e00" <= char <= "\u9fff"))
    digits = sum(1 for char in value if char.isdigit())
    return cjk >= 2 or latin >= 3 or (digits >= 4 and len(value) <= 24)


def _decode_title_candidates(value: str) -> list[str]:
    raw = normalize_space(str(value or "").strip().strip("\"'"))
    if not raw:
        return []
    candidates: dict[str, None] = {raw: None}

    if PERCENT_ESCAPE_RE.search(raw):
        for encoding in ("utf-8", "gb18030", "gbk"):
            try:
                candidates[unquote(raw, encoding=encoding, errors="replace")] = None
            except Exception:
                pass

    for candidate in list(candidates):
        for codec in ("gb18030", "gbk", "latin1", "cp1252"):
            try:
                repaired = candidate.encode(codec).decode("utf-8")
            except Exception:
                continue
            candidates[repaired] = None

    return [normalize_space(item) for item in candidates if normalize_space(item)]


def _strip_document_title_noise(value: str) -> str:
    value = normalize_space(value)
    if not value:
        return ""
    value = value.rsplit("/", 1)[-1].rsplit("\\", 1)[-1]
    value = value.split("?", 1)[0].split("#", 1)[0]
    value = DOCUMENT_EXTENSION_RE.sub("", value)
    value = value.replace("_", " ")
    value = re.sub(r"\s*[-\u2013\u2014]\s*$", "", value)
    return normalize_space(value)


def _content_disposition_filename(value: str) -> str:
    if not value:
        return ""
    filename_star = re.search(r"filename\*\s*=\s*(?:UTF-8''|utf-8'')?([^;]+)", value)
    filename = filename_star or re.search(r'filename\s*=\s*"?([^";]+)"?', value)
    return filename.group(1).strip() if filename else ""


def _is_opaque_document_title(value: str) -> bool:
    stem = DOCUMENT_EXTENSION_RE.sub("", normalize_space(value)).strip().lower()
    compact = re.sub(r"[\s._-]+", "", stem)
    if not compact:
        return True
    if _cjk_char_count(stem):
        return False
    if UUIDISH_RE.fullmatch(stem):
        return True
    if compact.isdigit() and len(compact) >= 8:
        return True
    if len(compact) >= 16 and all(char in "0123456789abcdef" for char in compact):
        return True
    latin = sum(char.isalpha() for char in stem)
    digits = sum(char.isdigit() for char in stem)
    return latin == 0 and digits >= max(8, len(compact) * 0.7)


def _title_score(value: str, source_weight: int = 0) -> float:
    cleaned = _strip_document_title_noise(value)
    if not cleaned or not _has_title_signal(cleaned):
        return -10_000.0
    cjk = _cjk_char_count(cleaned)
    latin = sum(1 for char in cleaned if char.isalpha() and not ("\u4e00" <= char <= "\u9fff"))
    digits = sum(1 for char in cleaned if char.isdigit())
    score = float(source_weight)
    score += min(cjk, 32) * 4.0
    score += min(latin, 40) * 0.8
    score += min(digits, 16) * 0.25
    score += min(len(cleaned), 80) * 0.08
    score -= sum(cleaned.count(marker) for marker in MOJIBAKE_MARKERS) * 30.0
    score -= len(PERCENT_ESCAPE_RE.findall(cleaned)) * 10.0
    if _is_opaque_document_title(cleaned):
        return -10_000.0
    if len(cleaned) > 120:
        score -= (len(cleaned) - 120) * 0.5
    return score


def _best_title(candidates: Iterable[tuple[str, int]]) -> str:
    best = ""
    best_score = -10_000.0
    for raw, source_weight in candidates:
        for decoded in _decode_title_candidates(raw):
            cleaned = _strip_document_title_noise(decoded)
            score = _title_score(cleaned, source_weight)
            if score > best_score:
                best = cleaned
                best_score = score
    return best if best_score > -9_000.0 else ""


def _document_text_title(text: str, limit: int = 96) -> str:
    cleaned = clean_document_text(text)
    if not cleaned:
        return ""
    parts = re.split(r"[\r\n\u3002\uff1b;]", cleaned, maxsplit=1)
    title = parts[0] if parts else cleaned
    title = normalize_space(title[:limit])
    return title if _has_title_signal(title) else ""


def normalize_document_title(
    url: str,
    filetype: str,
    text: str = "",
    fallback_title: str = "",
    content_disposition: str = "",
) -> str:
    """Return a searchable display title for downloaded documents."""

    try:
        parsed_path = urlparse(url or "").path
    except ValueError:
        parsed_path = url or ""

    candidates: list[tuple[str, int]] = []
    disposition_name = _content_disposition_filename(content_disposition)
    if disposition_name:
        candidates.append((disposition_name, 45))
    if parsed_path:
        candidates.append((parsed_path.rsplit("/", 1)[-1], 35))
    if fallback_title:
        candidates.append((fallback_title, 30))
    text_title = _document_text_title(text)
    if text_title:
        candidates.append((text_title, 10))

    best = _best_title(candidates)
    if best:
        return best
    suffix = (filetype or "document").lower().lstrip(".")
    return DOCUMENT_FALLBACK_TITLES.get(suffix, f"{suffix.upper()} document")


def _extract_document_metadata_title(content: bytes, filetype: str) -> str:
    filetype = filetype.lower()
    if filetype == "pdf":
        try:
            import pdfplumber

            with pdfplumber.open(io.BytesIO(content)) as pdf:
                metadata = pdf.metadata or {}
                return normalize_space(str(metadata.get("Title") or metadata.get("title") or ""))
        except Exception:
            return ""
    if filetype == "docx":
        try:
            from docx import Document

            document = Document(io.BytesIO(content))
            return normalize_space(str(document.core_properties.title or ""))
        except Exception:
            return ""
    if filetype == "pptx":
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                with archive.open("docProps/core.xml") as file:
                    return _xml_text(file.read(), tags={"title"})
        except Exception:
            return ""
    return ""


def extract_document_title(
    content: bytes,
    filetype: str,
    url: str,
    text: str = "",
    content_disposition: str = "",
    fallback_title: str = "",
) -> str:
    metadata_title = _extract_document_metadata_title(content, filetype)
    candidates = [(metadata_title, 55)] if metadata_title else []
    normalized = normalize_document_title(url, filetype, text, fallback_title, content_disposition)
    candidates.append((normalized, 40))
    return _best_title(candidates) or normalized


def _local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1].lower()


def _xml_text(payload: bytes, tags: set[str] | None = None) -> str:
    import xml.etree.ElementTree as ElementTree

    try:
        root = ElementTree.fromstring(payload)
    except Exception:
        return ""
    wanted = {tag.lower() for tag in (tags or {"t"})}
    values: list[str] = []
    for element in root.iter():
        if _local_name(element.tag) in wanted and element.text:
            values.append(element.text)
    return normalize_space(" ".join(values))


def _extract_pptx_text(content: bytes) -> str:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            slide_names = sorted(
                name
                for name in archive.namelist()
                if name.startswith(("ppt/slides/slide", "ppt/notesSlides/notesSlide")) and name.endswith(".xml")
            )
            values: list[str] = []
            for name in slide_names:
                with archive.open(name) as file:
                    text = _xml_text(file.read(), tags={"t"})
                if text:
                    values.append(text)
            return normalize_space(" ".join(values))
    except Exception:
        return ""


def _extract_legacy_ppt_text(content: bytes) -> str:
    """Binary .ppt needs an OLE-aware parser; byte-guessing creates false text."""

    return ""


def tokenize(text: str) -> list[str]:
    """Tokenize Chinese and Latin text for local TF-IDF tests and fallback search."""

    if not text:
        return []
    tokens: list[str] = []
    try:
        import jieba

        tokens.extend(token.strip().lower() for token in jieba.lcut(text) if token.strip())
    except Exception:
        for match in TOKEN_RE.finditer(text):
            token = match.group(0).lower()
            if not CJK_RE.fullmatch(token):
                tokens.append(token)
            else:
                tokens.extend(_cjk_ngrams(token))

    for cjk in CJK_RE.findall(text):
        tokens.extend(_cjk_ngrams(cjk))
    return list(dict.fromkeys(token for token in tokens if TOKEN_RE.search(token)))


def _cjk_ngrams(value: str) -> list[str]:
    if not value:
        return []
    grams = [value] if len(value) <= 8 else []
    for size in (1, 2, 3, 4):
        if len(value) >= size:
            grams.extend(value[index : index + size] for index in range(0, len(value) - size + 1))
    return grams


def term_frequency(text: str) -> Counter[str]:
    return Counter(tokenize(text))


def cosine_similarity(query: str, document: str, idf: dict[str, float] | None = None) -> float:
    query_tf = term_frequency(query)
    doc_tf = term_frequency(document)
    if not query_tf or not doc_tf:
        return 0.0
    idf = idf or {}
    terms = set(query_tf) | set(doc_tf)
    dot = 0.0
    query_norm = 0.0
    doc_norm = 0.0
    for term in terms:
        weight = idf.get(term, 1.0)
        q = query_tf.get(term, 0) * weight
        d = doc_tf.get(term, 0) * weight
        dot += q * d
        query_norm += q * q
        doc_norm += d * d
    if query_norm == 0.0 or doc_norm == 0.0:
        return 0.0
    return dot / (math.sqrt(query_norm) * math.sqrt(doc_norm))


def build_idf(documents: Iterable[str]) -> dict[str, float]:
    docs = list(documents)
    total = len(docs)
    df: Counter[str] = Counter()
    for doc in docs:
        df.update(set(tokenize(doc)))
    return {term: math.log((1 + total) / (1 + count)) + 1.0 for term, count in df.items()}


def html_to_text(html: str) -> tuple[str, str, list[str]]:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html or "", "lxml")
    for element in soup(["script", "style", "noscript"]):
        element.decompose()
    title = normalize_space(soup.title.get_text(" ")) if soup.title else ""
    anchors = [normalize_space(a.get_text(" ")) for a in soup.find_all("a")]
    anchors = [item for item in anchors if item]
    text = normalize_space(soup.get_text(" "))
    return title, text, anchors


STATIC_ASSET_EXTENSIONS = {
    "avif",
    "bmp",
    "css",
    "gif",
    "ico",
    "jpeg",
    "jpg",
    "js",
    "m4a",
    "mkv",
    "mov",
    "mp3",
    "mp4",
    "ogg",
    "otf",
    "png",
    "svg",
    "ttf",
    "wav",
    "webm",
    "webp",
    "woff",
    "woff2",
}
CRAWLABLE_EXTENSIONS = {
    "",
    "asp",
    "aspx",
    "doc",
    "docx",
    "htm",
    "html",
    "jsp",
    "pdf",
    "php",
    "ppt",
    "pptx",
    "shtml",
    "shtm",
    "txt",
    "xls",
    "xlsx",
    "xml",
}
LINK_ATTRIBUTES = (
    "href",
    "src",
    "action",
    "data-url",
    "data-href",
    "data-src",
    "data-original",
    "data-link",
    "data-path",
)
ABSOLUTE_URL_RE = re.compile(r"https?://[^\s\"'<>\\)]+", re.IGNORECASE)
ROOT_PATH_RE = re.compile(r"(?<![A-Za-z0-9_])/(?:[A-Za-z0-9._~!$&'()*+,;=:@%/-]+)(?:\?[A-Za-z0-9._~!$&'()*+,;=:@%/?-]*)?")
RESOURCE_DB_RE = re.compile(r"/openlist/d/resource/(anime|comic|game|novel)/([acgn]\d{3,})/", re.IGNORECASE)
DBID_RE = re.compile(r"\\?\"dbId\\?\"\s*:\s*\\?\"([acgn]\d{3,})\\?\"", re.IGNORECASE)
DB_PREFIX_ROUTES = {"a": "anime", "c": "comic", "g": "game", "n": "novel"}


def extract_12club_resource_links(payload: str) -> list[str]:
    links: dict[str, None] = {}
    for route, db_id in RESOURCE_DB_RE.findall(payload or ""):
        links[f"http://12club.nankai.edu.cn/{route.lower()}/{db_id.lower()}"] = None
    for db_id in DBID_RE.findall(payload or ""):
        db_id = db_id.lower()
        route = DB_PREFIX_ROUTES.get(db_id[0])
        if route:
            links[f"http://12club.nankai.edu.cn/{route}/{db_id}"] = None
    return list(links)


def extract_links(html: str, base_url: str) -> list[str]:
    """Extract crawlable links from HTML, inline app data, and sitemap-like text.

    Many NKU sites are WebPlus/Sudy pages and 12club is a Next.js app. A-tag
    extraction alone misses URLs embedded in data attributes, inline RSC data,
    sitemap XML, and resource paths. This keeps the crawler page-focused by
    rejecting static assets while still turning resource IDs into canonical
    detail-page URLs.
    """

    from urllib.parse import urldefrag, urljoin, urlparse
    from bs4 import BeautifulSoup, XMLParsedAsHTMLWarning

    parser = "xml" if _looks_like_xml(html) else "lxml"
    with warnings.catch_warnings():
        warnings.filterwarnings("ignore", category=XMLParsedAsHTMLWarning)
        soup = BeautifulSoup(html or "", parser)
    links: list[str] = []

    def add_candidate(value: object) -> None:
        if value is None:
            return
        if isinstance(value, (list, tuple)):
            for item in value:
                add_candidate(item)
            return
        href = str(value).strip()
        if not href:
            return
        lowered = href.lower()
        if lowered.startswith(("javascript:", "mailto:", "tel:", "data:", "#")):
            return
        try:
            absolute = urldefrag(urljoin(base_url, href))[0]
            parsed = urlparse(absolute)
        except ValueError:
            return
        if parsed.scheme in {"http", "https"} and parsed.netloc and _is_crawlable_url(absolute):
            links.append(absolute)

    for element in soup.find_all(True):
        for attribute in LINK_ATTRIBUTES:
            add_candidate(element.get(attribute))
        for attribute, value in element.attrs.items():
            if attribute.startswith("data-"):
                add_candidate(value)

    source = html or ""
    for match in ABSOLUTE_URL_RE.finditer(source):
        add_candidate(match.group(0).rstrip('.,;'))
    for match in ROOT_PATH_RE.finditer(source):
        path = match.group(0)
        if match.start() > 0 and source[match.start() - 1] == "<":
            continue
        if path.startswith(("//", "/_next/static/", "/static/")):
            continue
        add_candidate(path.rstrip('.,;'))

    for resource_link in extract_12club_resource_links(source):
        add_candidate(resource_link)

    return list(dict.fromkeys(links))


def extract_asset_links(html: str, base_url: str, extensions: set[str] | None = None) -> list[str]:
    """Extract static asset URLs when crawler logic needs to inspect app bundles."""

    from urllib.parse import urldefrag, urljoin, urlparse
    from bs4 import BeautifulSoup

    wanted = {item.lower().lstrip(".") for item in (extensions or {"js"})}
    soup = BeautifulSoup(html or "", "lxml")
    links: list[str] = []
    for element in soup.find_all(True):
        for attribute in ("src", "href"):
            value = element.get(attribute)
            if not value:
                continue
            try:
                absolute = urldefrag(urljoin(base_url, str(value).strip()))[0]
                parsed = urlparse(absolute)
            except ValueError:
                continue
            suffix = parsed.path.rsplit(".", 1)[-1].lower() if "." in parsed.path.rsplit("/", 1)[-1] else ""
            if parsed.scheme in {"http", "https"} and parsed.netloc and suffix in wanted:
                links.append(absolute)
    return list(dict.fromkeys(links))


def _looks_like_xml(value: str) -> bool:
    prefix = (value or "").lstrip()[:120].lower()
    return prefix.startswith("<?xml") or prefix.startswith("<urlset") or prefix.startswith("<sitemapindex") or prefix.startswith("<rss") or prefix.startswith("<feed")


def _is_crawlable_url(url: str) -> bool:
    from urllib.parse import urlparse

    try:
        parsed = urlparse(url)
    except ValueError:
        return False
    path = parsed.path.lower()
    filename = path.rsplit("/", 1)[-1]
    suffix = filename.rsplit(".", 1)[-1] if "." in filename else ""
    if suffix in STATIC_ASSET_EXTENSIONS:
        return False
    if suffix and suffix not in CRAWLABLE_EXTENSIONS:
        return False
    return True


def infer_filetype(url: str, content_type: str = "") -> str:
    suffix = url.split("?", 1)[0].rsplit(".", 1)
    if len(suffix) == 2 and 1 <= len(suffix[1]) <= 5:
        return suffix[1].lower()
    if "pdf" in content_type:
        return "pdf"
    if "word" in content_type:
        return "docx"
    if "excel" in content_type or "spreadsheet" in content_type:
        return "xlsx"
    return "html"


def extract_document_text(content: bytes, filetype: str) -> str:
    filetype = filetype.lower()
    if filetype == "pdf":
        try:
            import pdfplumber

            with pdfplumber.open(io.BytesIO(content)) as pdf:
                return normalize_space(" ".join((page.extract_text() or "") for page in pdf.pages))
        except Exception:
            return ""
    if filetype in {"docx", "doc"}:
        if filetype == "doc":
            return ""
        try:
            from docx import Document

            document = Document(io.BytesIO(content))
            return normalize_space(" ".join(paragraph.text for paragraph in document.paragraphs))
        except Exception:
            return ""
    if filetype in {"xlsx", "xls"}:
        if filetype == "xls":
            return ""
        try:
            from openpyxl import load_workbook

            with warnings.catch_warnings():
                warnings.filterwarnings(
                    "ignore",
                    message="Print area cannot be set to Defined name:.*",
                    category=UserWarning,
                    module="openpyxl.*",
                )
                workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            values: list[str] = []
            for sheet in workbook.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    values.extend(str(cell) for cell in row if cell is not None)
            return normalize_space(" ".join(values))
        except Exception:
            return ""
    if filetype in {"pptx", "ppt"}:
        if filetype == "pptx":
            return _extract_pptx_text(content)
        return _extract_legacy_ppt_text(content)
    try:
        return clean_document_text(content.decode("utf-8", errors="ignore"), filetype)
    except Exception:
        return ""


def make_snippet(text: str, query: str, width: int = 180) -> str:
    text = normalize_space(text)
    if not text:
        return ""
    terms = tokenize(query)
    positions = [text.lower().find(term.lower()) for term in terms if term and text.lower().find(term.lower()) >= 0]
    start = max(min(positions) - width // 3, 0) if positions else 0
    snippet = text[start : start + width]
    return snippet + ("..." if start + width < len(text) else "")

